#!/usr/bin/env python3
"""MCQ Parser — .docx or .html → standard MCQ Markdown."""
import sys, os, re

INPUT = sys.argv[1]
OUTDIR = sys.argv[2] if len(sys.argv) > 2 else os.path.join(os.path.dirname(INPUT), 'parsed')
EXT = INPUT.rsplit('.', 1)[-1].lower()

items = []

# ── Input: .docx ──
if EXT == 'docx':
    from docx import Document
    from docx.oxml.ns import qn
    doc = Document(INPUT)

    # Extract all images to OUTDIR/images
    img_dir = os.path.join(OUTDIR, 'images')
    os.makedirs(img_dir, exist_ok=True)
    rel_map = {}  # rId -> filename
    for rel_id, rel in doc.part.rels.items():
        if 'image' in rel.reltype:
            ext = rel.target_ref.split('.')[-1]
            fname = f"img_{rel_id}.{ext}"
            with open(os.path.join(img_dir, fname), 'wb') as f:
                f.write(rel.target_part.blob)
            rel_map[rel_id] = fname

    pi = ti = 0
    for child in doc.element.body:
        if child.tag == qn('w:p'):
            # Check for inline images in this paragraph
            has_img = False
            for run in doc.paragraphs[pi].runs:
                for blip in run._element.findall('.//' + qn('a:blip')):
                    embed = blip.get(qn('r:embed'))
                    if embed and embed in rel_map:
                        items.append(('img', f'images/{rel_map[embed]}'))
                        has_img = True
            t = doc.paragraphs[pi].text.strip()
            pi += 1
            if t:
                items.append(('p', t))
            elif has_img:
                pass  # image-only paragraph, already added
        elif child.tag == qn('w:tbl') and ti < len(doc.tables):
            tbl = doc.tables[ti]
            ti += 1
            rows = []
            for row in tbl.rows:
                cells = [cell.text.strip() for cell in row.cells]
                while cells and cells[-1] == '': cells.pop()
                while cells and cells[0] == '': cells.pop(0)
                if cells: rows.append(cells)
            # Detect ABCD option table (first column of data rows = A/B/C/D)
            is_opt = len(rows) >= 5 and all(rows[i][0] in 'ABCD' for i in range(1, min(5, len(rows))))
            md_rows = []
            for ri, cells in enumerate(rows):
                if is_opt and ri == 0: continue  # skip header for option tables
                md_rows.append('| ' + ' | '.join(cells) + ' |')
            if len(md_rows) >= 2:
                md = md_rows[0] + '\n|' + '|'.join(['---'] * len(md_rows[0].split('|'))) + '|'
                for r in md_rows[1:]:
                    md += '\n' + r
                items.append(('table', md))

# ── Input: .html ──
elif EXT == 'html':
    from html.parser import HTMLParser

    class MCQHTMLParser(HTMLParser):
        def __init__(self):
            super().__init__()
            self.items = []
            self.cur = ''
            self.ip = False
            self.td = 0
            self.trs = []
            self.tcs = []
            self.ir = False

        def handle_starttag(self, tag, attrs):
            if tag == 'p':
                self.ip = True
                self.cur = ''
            elif tag == 'img':
                for k, v in attrs:
                    if k == 'src':
                        self.items.append(('img', v))
            elif tag == 'table':
                self.td += 1
                self.trs = []
            elif tag == 'tr' and self.td > 0:
                self.ir = True
                self.tcs = []
            elif tag in ('td', 'th') and self.ir:
                self.cur = ''

        def handle_endtag(self, tag):
            if tag == 'p':
                self.ip = False
                t = self.cur.strip()
                if t:
                    self.items.append(('p', t))
            elif tag == 'table':
                self.td -= 1
                if self.td == 0 and self.trs:
                    ls = []
                    for ri, r in enumerate(self.trs):
                        ls.append('| ' + ' | '.join(r) + ' |')
                        if ri == 0 and r:
                            ls.append('|' + '|'.join(['---'] * len(r)) + '|')
                    self.items.append(('table', '\n'.join(ls)))
            elif tag == 'tr':
                self.ir = False
                if self.tcs:
                    self.trs.append(self.tcs)
            elif tag in ('td', 'th'):
                self.tcs.append(self.cur.strip())

        def handle_data(self, data):
            if self.ip:
                self.cur += data
            elif self.ir:
                self.cur += data

    with open(INPUT, encoding='utf-8') as f:
        html = f.read()
    p = MCQHTMLParser()
    p.feed(html)
    items = p.items

else:
    print(f"Unsupported format: .{EXT}")
    sys.exit(1)

# ── Find first question ──
start = 0
for i, (tag, text) in enumerate(items):
    if tag == 'p' and re.match(r'^\d{1,2}\s+[A-Z]', text):
        start = i
        break
items = items[start:]

# ── Group by question number ──
blocks, cur = [], []
for el in items:
    tag, text = el
    if tag == 'p' and re.match(r'^\d{1,2}\s+[A-Z]', text) and cur:
        blocks.append(cur)
        cur = [el]
    else:
        cur.append(el)
if cur:
    blocks.append(cur)

# ── Parse each block ──
questions = []
for block in blocks:
    if not block or block[0][0] != 'p':
        continue
    first_text = block[0][1]

    # Fix 3: Skip copyright block
    if 'Permission to reproduce' in first_text:
        continue

    m = re.match(r'^(\d{1,2})\s+', first_text)
    if not m:
        continue
    num = int(m.group(1))
    if not (1 <= num <= 30):
        continue

    qbody = [re.sub(r'^\d{1,2}\s+', '', first_text)]
    opts = []
    images_in_q = []
    tables_in_q = []

    for tag, text in block[1:]:
        # Fix 3: copyright text in block → skip
        if tag == 'p' and 'Permission to reproduce' in text:
            break
        if tag == 'img':
            images_in_q.append(f'![image]({text})')
        elif tag == 'table':
            tables_in_q.append(f'\n{text}\n')
            # Also extract options from table markdown if first column is ABCD
            tbl_rows = [r for r in text.split('\n') if r.strip() and not r.startswith('|---')]
            if len(tbl_rows) >= 4:
                first_col = []
                for r in tbl_rows:
                    parts = [c.strip() for c in r.split('|') if c.strip()]
                    if parts: first_col.append(parts[0])
                if len(first_col) >= 4 and all(v in 'ABCD' for v in first_col[:4]):
                    for ri in range(min(4, len(tbl_rows))):
                        parts = [c.strip() for c in tbl_rows[ri].split('|') if c.strip()]
                        if parts:
                            label = parts[0]
                            rest = ' | '.join(parts[1:]) if len(parts) > 1 else ''
                            if label in 'ABCD' and label not in {o[0] for o in opts}:
                                opts.append((label, '—'))
        elif tag == 'p':
            t = text.strip()
            if not t or re.match(r'^\d+$', t):
                continue
            # Fix 3: paper code line at end
            if re.match(r'^\d{4}/\d{2}/', t):
                continue

            # Detect LaTeX table: @lll@ & & year 1 & year 2 \\ agriculture & 4 & 3 \\
            if re.match(r'^@[lcr| ]+@', t) or ('&' in t and '\\\\' in t):
                # Convert LaTeX tabular to markdown table
                latex = t.replace('\\\\', '\n').strip()
                # Remove column spec like @lll@
                latex = re.sub(r'^@[^@]*@\s*', '', latex)
                rows = [r.strip() for r in latex.split('\n') if r.strip()]
                md_rows = []
                for ri, row in enumerate(rows):
                    cells = [c.strip() for c in row.split('&')]
                    md_rows.append('| ' + ' | '.join(cells) + ' |')
                    if ri == 0 and cells:
                        md_rows.append('|' + '|'.join(['---'] * len(cells)) + '|')
                if md_rows:
                    tables_in_q.append('\n' + '\n'.join(md_rows) + '\n')
                continue

            m2 = re.match(r'^([A-D])\s+(.+)', t)
            if m2:
                opts.append((m2.group(1), m2.group(2)))
            else:
                qbody.append(t)

    # Fix 1: unlabeled text before B/C/D → it's the missing A
    labels_found = {o[0] for o in opts}
    if 'A' not in labels_found and labels_found and qbody:
        # Last item in qbody that is not an image/table marker is likely the missing A
        text_items = [x for x in qbody if not x.startswith('![') and not x.startswith('\n|')]
        if text_items:
            candidate = text_items[-1]
            # Only if it's short (option-like) and not clearly question text
            if len(candidate) < 200:
                qbody.remove(candidate)
                opts.insert(0, ('A', candidate))

    # Fix 3: inline options "A 0.8 B 0.9 C 1.2 D 1.25" or "0.8 B 0.9 C 1.2 D 1.25"
    expanded_opts = []
    for label, text in opts:
        # Check if this option text contains B/C/D labels inline
        if re.search(r'\s+[B-D]\s+', text):
            parts = re.split(r'\s+(?=[B-D]\s)', text)
            expanded_opts.append((label, parts[0].strip()))
            for part in parts[1:]:
                m3 = re.match(r'^([B-D])\s+(.+)', part)
                if m3:
                    expanded_opts.append((m3.group(1), m3.group(2).strip()))
        else:
            expanded_opts.append((label, text))
    opts = expanded_opts
    del expanded_opts

    # Fix 3b: inline options in qbody "0.8 B 0.9 C 1.2 D 1.25" (no A label at all)
    for idx, item in enumerate(qbody):
        if re.search(r'\s+B\s+.*\s+C\s+.*\s+D\s+', item):
            parts = re.split(r'\s+(?=[BCD]\s)', item)
            if len(parts) >= 3:
                # First part before B = option A (may be empty if text starts with B)
                a_part = parts[0].strip()
                if 'A' not in labels_found:
                    opts.insert(0, ('A', a_part if a_part else '⚠️ Missing in source'))
                # B, C, D parts
                for part in parts[1:]:
                    m3 = re.match(r'^([B-D])\s+(.+)', part)
                    if m3 and m3.group(1) not in labels_found:
                        opts.append((m3.group(1), m3.group(2)))
                qbody[idx] = ''  # Remove from question text
    qbody = [x for x in qbody if x]

    # Fix 4: image-only question
    if not opts and images_in_q:
        opts = [('A', 'See diagram'), ('B', 'See diagram'),
                ('C', 'See diagram'), ('D', 'See diagram')]

    # Assemble question text
    q_parts = qbody + images_in_q + tables_in_q
    qt = ' '.join(q_parts).strip()

    # Fill remaining slots: only from qbody if no tables present (tables mean options are visual)
    if not tables_in_q:
        while len(opts) < 4:
            next_label = chr(65 + len(opts))
            text_items = [x for x in qbody if not x.startswith('![') and len(x) > 3]
            if text_items:
                opts.append((next_label, text_items[-1]))
                qbody = [x for x in qbody if x != text_items[-1]]
            else:
                opts.append((next_label, '⚠️ Missing'))
    else:
        while len(opts) < 4:
            opts.append((chr(65 + len(opts)), '⚠️ Missing'))
    questions.append((num, qt, opts[:4]))

# ── Group papers (Q1 = new paper) ──
papers, cp = [], []
for q in questions:
    if q[0] == 1 and cp:
        papers.append(cp)
        cp = []
    cp.append(q)
if cp:
    papers.append(cp)

# ── Save ──
os.makedirs(OUTDIR, exist_ok=True)
KNOWN = [
    ('0455/11', 'May/June', '2021'), ('0455/12', 'May/June', '2021'),
    ('0455/13', 'May/June', '2021'), ('0455/11', 'Oct/Nov', '2021'),
    ('0455/12', 'Oct/Nov', '2021'), ('0455/13', 'Oct/Nov', '2021'),
    ('0455/11', 'May/June', '2022'), ('0455/12', 'May/June', '2022'),
    ('0455/13', 'May/June', '2022'), ('0455/11', 'Oct/Nov', '2022'),
]

for pi, paper in enumerate(papers):
    if pi < len(KNOWN):
        code, session, year = KNOWN[pi]
    else:
        code, session, year = ('0455_XX', 'Unknown', f'P{pi+1}')
    fname = f"{code.replace('/','_')}_{session.replace('/','')}_{year}.md"
    path = os.path.join(OUTDIR, fname)
    complete = sum(1 for _, _, o in paper if o[3][1])

    with open(path, 'w', encoding='utf-8') as f:
        f.write(f"## {code} {session} {year} · IGCSE Economics Paper 1\n\n")
        letters = 'ABCD'
        for num, qt, opts in paper:
            missing = sum(1 for _, t in opts if not t)
            f.write(f"### Q{num}")
            if missing:
                f.write(f" ⚠️ missing {missing}")
            f.write(f"\n{qt}\n\n")
            for li, (_, t) in enumerate(opts):
                if t:
                    f.write(f"- {letters[li]}) {t}\n")
                else:
                    f.write(f"- {letters[li]}) ⚠️ MISSING\n")
            f.write(f"\n> Answer: ?\n\n")

    imgs = sum(1 for _, qt, _ in paper if '![' in qt)
    tbls = sum(1 for _, qt, _ in paper if '|---' in qt)
    print(f"  {fname}: {len(paper)} Qs ({complete} OK), {imgs} img, {tbls} tbl")

total = sum(len(p) for p in papers)
print(f"\nTotal: {total} questions in {len(papers)} papers")
print(f"Saved: {OUTDIR}")
